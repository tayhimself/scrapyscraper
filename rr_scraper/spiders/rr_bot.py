import scrapy
from htmldocx import HtmlToDocx
from docx import Document
from docx.shared import Inches
import requests

class RrBotSpider(scrapy.Spider):
    name = 'rr_bot'
    allowed_domains = ['railroadersleep.fra.dot.gov']
    start_urls = ['http://railroadersleep.fra.dot.gov/']

    def __init__(self):
        self.links=[]
        self.links.append('http://railroadersleep.fra.dot.gov/')
        self.document = Document()
        self.filename = f'rrsleep.docx'



    def parse(self, response):
        title = response.css("h1.cms-replaceable::text").get()
        # htmlfilename = f'{title}.html'

        document = Document()
        new_parser = HtmlToDocx()

        # get all divs with cms-editable class
        divs = response.css("div.cms-editable")
        aside = response.css("aside")
        # image titles
        img_titles = response.css("div.image-group h2").getall()

        # Add the divs to a string
        content = ""
        for div in divs:
            content += div.get()
        for a in aside:
            content += a.get()
        # Add the image titles to the string
        for i in range(len(img_titles)):
            content += img_titles[i]

        # get images
        images = response.css("div.image-group img").getall()

        # add the content to the document
        if content != "":
            new_parser.add_html_to_document(content, self.document)
            self.document.save(self.filename)

        # add the images to the document
        for i in range(len(images)):
            # get the image url
            img_url = images[i].split('src="')[1].split('"')[0]
            abs_img_url = response.urljoin(img_url)
            # get the image
            resp = requests.get(abs_img_url, stream=True)
            if resp.status_code == 200:
                with open('out.png', 'wb') as f:
                    for chunk in resp:
                        f.write(chunk)
                self.document.add_picture('out.png', width=Inches(6))
                self.document.save(self.filename)

        self.log(f'Saved file {self.filename}')


        for href in response.css('a::attr(href)'):
            self.links.append(response.url)
            yield response.follow(href, self.parse)
